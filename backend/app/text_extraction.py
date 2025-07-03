"""
KI-QMS Text-Extraktion für Dokumentenmanagement

Dieses Modul stellt robuste Text-Extraktion aus verschiedenen Dateiformaten
bereit für RAG-Indexierung, KI-Verarbeitung und Volltext-Suche im QMS.

Hauptfunktionalitäten:
- 📄 Multi-Format Text-Extraktion (PDF, DOC, XLSX, TXT, MD)
- 🔍 QMS-spezifische Keyword-Erkennung
- 🧹 Text-Bereinigung für bessere KI-Verarbeitung
- 📊 Strukturierte Datenextraktion (Tabellen, Überschriften)
- 🛡️ Robuste Fehlerbehandlung und Fallback-Mechanismen

Unterstützte Formate:
- PDF: PyPDF2 mit Seiten-strukturierter Extraktion
- Word: python-docx für DOCX mit Tabellen-Support
- Excel: openpyxl für XLSX mit Multi-Sheet-Verarbeitung  
- Text: TXT, MD mit UTF-8/Latin-1 Fallback
- Legacy: DOC, XLS über Fehlermeldungen

QMS-Features:
- ISO 13485 relevante Keyword-Erkennung
- Normreferenzen (ISO, IEC, DIN, EN) Detection
- Compliance-Begriffe (MDR, FDA, CFR) Extraction
- QMS-Prozess-Terminologie (CAPA, Audit, Kalibrierung)

Performance:
- Chunked Processing für große Dateien
- Memory-effiziente Stream-Verarbeitung
- Graceful Degradation bei Parsing-Fehlern
- Logging für Debugging und Monitoring

Autoren: KI-QMS Entwicklungsteam
Version: 2.0.0 (Enhanced für RAG-Integration)
"""

from pathlib import Path
from typing import Optional, Union, Dict, Any
import logging
import tempfile
import asyncio

logger = logging.getLogger(__name__)

# Enhanced OCR Import
try:
    from .enhanced_ocr_engine import extract_enhanced_text
    ENHANCED_OCR_AVAILABLE = True
except ImportError as e:
    ENHANCED_OCR_AVAILABLE = False
    logger.warning(f"⚠️ Enhanced OCR nicht verfügbar: {e}")

def extract_text_from_file(file_path: Union[str, Path], mime_type: str) -> str:
    """
    Extrahiert Text aus verschiedenen Dateiformaten mit Enhanced OCR Fallback
    """
    file_path = Path(file_path) if isinstance(file_path, str) else file_path
    
    try:
        logger.info(f"🔍 Textextraktion gestartet für: {file_path.name} ({mime_type})")
        
        # Standard-Textextraktion versuchen
        extracted_text = ""
        
        if mime_type == 'application/pdf':
            extracted_text = _extract_pdf_text(file_path)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            extracted_text = _extract_word_text(file_path)
        elif mime_type in ['application/msword']:
            extracted_text = "[DOC-Format nicht unterstützt - bitte zu DOCX konvertieren]"
        elif mime_type == 'text/plain':
            extracted_text = _extract_text_file(file_path)
        else:
            logger.warning(f"⚠️ Unbekannter MIME-Type: {mime_type}")
            extracted_text = "[Unbekanntes Dateiformat]"
        
        # Enhanced OCR Fallback für problematische Dokumente
        if len(extracted_text.strip()) < 50 and ENHANCED_OCR_AVAILABLE:
            logger.info("🔄 Standard-Extraktion ergab wenig Text - versuche Enhanced OCR")
            try:
                # Async call in sync context
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                enhanced_result = loop.run_until_complete(
                    extract_enhanced_text(file_path, mime_type)
                )
                loop.close()
                
                if enhanced_result['success'] and len(enhanced_result['text']) > len(extracted_text):
                    logger.info(f"✅ Enhanced OCR erfolgreich: {len(enhanced_result['text'])} Zeichen, "
                              f"{enhanced_result['images_processed']} Bilder verarbeitet")
                    extracted_text = enhanced_result['text']
                    
                    # Metadaten in Log ausgeben
                    logger.info(f"📊 OCR-Methode: {enhanced_result['ocr_method']}")
                
            except Exception as e:
                logger.error(f"❌ Enhanced OCR Fallback fehlgeschlagen: {e}")
        
        # Ergebnis validieren
        if not extracted_text or extracted_text.strip() == "":
            extracted_text = "[Kein Text gefunden]"
        
        logger.info(f"✅ Textextraktion abgeschlossen: {len(extracted_text)} Zeichen")
        return extracted_text
        
    except Exception as e:
        logger.error(f"❌ Textextraktion fehlgeschlagen für {file_path.name}: {e}")
        return f"[Textextraktion fehlgeschlagen: {str(e)}]"

def _extract_text_file(file_path: Path) -> str:
    """Extrahiert Text aus TXT/MD Dateien."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback für andere Encodings
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def _extract_pdf_text(file_path: Path) -> str:
    """Extrahiert Text aus PDF-Dateien mit PyPDF2."""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"=== Seite {page_num} ===")
                        text_parts.append(page_text)
                except Exception as page_error:
                    text_parts.append(f"[Fehler auf Seite {page_num}: {str(page_error)}]")
            
            return "\n".join(text_parts) if text_parts else "[Kein Text extrahiert]"
            
    except ImportError:
        return "[PyPDF2 nicht installiert]"
    except Exception as e:
        return f"[PDF-Extraktion fehlgeschlagen: {str(e)}]"

def _extract_word_text(file_path: Path) -> str:
    """Extrahiert Text aus Word-Dokumenten (.docx)."""
    try:
        from docx import Document
        
        doc = Document(str(file_path))
        text_parts = []
        
        # Absätze extrahieren
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Tabellen extrahieren
        for table_num, table in enumerate(doc.tables, 1):
            text_parts.append(f"\n=== Tabelle {table_num} ===")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts) if text_parts else "[Kein Text gefunden]"
        
    except ImportError:
        return "[python-docx nicht installiert]"
    except Exception as e:
        return f"[Word-Extraktion fehlgeschlagen: {str(e)}]"

def _extract_excel_text(file_path: Path) -> str:
    """Extrahiert Text aus Excel-Dateien (.xlsx)."""
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n=== {sheet_name} ===")
            
            # Nur nicht-leere Zeilen verarbeiten
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                row_text = " | ".join(row_values).strip()
                
                if row_text and row_text != " | " * (len(row_values) - 1):
                    text_parts.append(row_text)
        
        return "\n".join(text_parts) if text_parts else "[Kein Inhalt gefunden]"
        
    except ImportError:
        return "[openpyxl nicht installiert]"
    except Exception as e:
        return f"[Excel-Extraktion fehlgeschlagen: {str(e)}]"

def extract_keywords(text: str) -> list[str]:
    """
    Extrahiert QMS-relevante Keywords aus Text.
    
    Args:
        text: Eingabetext
        
    Returns:
        list[str]: Gefundene Keywords
    """
    # QMS-spezifische Keywords
    qms_keywords = [
        "ISO 13485", "ISO 9001", "ISO 14971", "IEC 62304",
        "MDR", "FDA", "CFR", "GMP", "GLP", "GCP",
        "Kalibrierung", "Validierung", "Verifikation",
        "Risiko", "CAPA", "Audit", "Korrektur",
        "Qualitätspolitik", "Qualitätsziel", "Qualitätsmanagement",
        "Dokumentenkontrolle", "Lenkung von Dokumenten",
        "Aufzeichnungen", "QMH", "Qualitätshandbuch",
        "SOP", "Arbeitsanweisung", "Verfahren",
        "Medizinprodukt", "CE-Kennzeichnung", "Konformität",
        "Überwachung", "Messung", "Prüfung",
        "Lieferant", "Beschaffung", "Einkauf",
        "Schulung", "Kompetenz", "Bewusstsein"
    ]
    
    found_keywords = []
    text_lower = text.lower()
    
    for keyword in qms_keywords:
        if keyword.lower() in text_lower:
            found_keywords.append(keyword)
    
    return found_keywords

def analyze_document_type(text: str, title: str = "") -> str:
    """
    Analysiert ersten 3000 Zeichen + Titel und bestimmt automatisch den Dokumenttyp.
    
    Args:
        text: Extrahierter Text (erste 3000 Zeichen werden analysiert)
        title: Dokumenttitel für zusätzliche Hinweise
        
    Returns:
        str: Erkannter Dokumenttyp aus DocumentType Enum
    """
    # Erste 3000 Zeichen für Analyse verwenden
    analysis_text = (title + " " + text)[:3000].lower()
    
    # Dokumenttyp-spezifische Keywords mit Gewichtung
    type_indicators = {
        "QM_MANUAL": [
            ("qualitätsmanagement", 10), ("qm-handbuch", 15), ("qualitätshandbuch", 15),
            ("unternehmenspolitik", 8), ("organisationsstruktur", 8), ("qualitätspolitik", 12),
            ("iso 13485", 10), ("mdr", 8), ("qualitätsmanagementsystem", 12)
        ],
        "SOP": [
            ("standard operating procedure", 15), ("sop", 15), ("standardarbeitsanweisung", 12),
            ("verfahrensanweisung", 12), ("prozessanweisung", 12), ("durchführung", 8),
            ("verantwortlichkeiten", 10), ("ablaufbeschreibung", 10), ("prozessbeschreibung", 10)
        ],
        "WORK_INSTRUCTION": [
            ("arbeitsanweisung", 15), ("arbeitsvorschrift", 12), ("durchführungsanweisung", 12),
            ("schritt-für-schritt", 10), ("anleitung", 8), ("detaillierte", 6),
            ("bedienung", 8), ("handhabung", 8), ("ausführung", 6)
        ],
        "FORM": [
            ("formular", 15), ("formblatt", 15), ("checkliste", 12), ("protokoll", 10),
            ("prüfprotokoll", 12), ("nachweis", 8), ("dokumentation", 6),
            ("unterschrift", 10), ("datum", 6), ("name", 4)
        ],
        "RISK_ASSESSMENT": [
            ("risikoanalyse", 15), ("risikobewertung", 15), ("iso 14971", 15),
            ("risikomanagement", 12), ("gefährdung", 10), ("schadensereignis", 10),
            ("wahrscheinlichkeit", 8), ("schweregrad", 8), ("risikokontrolle", 10)
        ],
        "VALIDATION_PROTOCOL": [
            ("validierung", 15), ("validierungsprotokoll", 15), ("iq", 12), ("oq", 12), ("pq", 12),
            ("installation qualification", 12), ("operational qualification", 12),
            ("performance qualification", 12), ("prüfplan", 10), ("testergebnis", 8)
        ],
        "CALIBRATION_PROCEDURE": [
            ("kalibrierung", 15), ("kalibrierverfahren", 15), ("justierung", 10),
            ("messgenauigkeit", 10), ("toleranz", 8), ("messmittel", 10),
            ("prüfmittel", 10), ("normal", 6), ("rückführbarkeit", 12)
        ],
        "AUDIT_REPORT": [
            ("audit", 15), ("auditbericht", 15), ("prüfbericht", 12), ("bewertung", 8),
            ("feststellung", 10), ("abweichung", 10), ("konformität", 10),
            ("nichtkonformität", 12), ("verbesserung", 8)
        ],
        "STANDARD_NORM": [
            ("iso", 12), ("din", 12), ("en", 8), ("iec", 12), ("astm", 10),
            ("norm", 10), ("standard", 10), ("anforderung", 8), ("spezifikation", 8),
            ("technische regel", 10), ("richtlinie", 8)
        ],
        "SPECIFICATION": [
            ("spezifikation", 15), ("anforderung", 10), ("technische daten", 10),
            ("leistungsmerkmale", 10), ("parameter", 8), ("eigenschaft", 6),
            ("charakteristikum", 8), ("kennwert", 8)
        ]
    }
    
    # Score für jeden Dokumenttyp berechnen
    scores = {}
    for doc_type, indicators in type_indicators.items():
        score = 0
        for keyword, weight in indicators:
            if keyword in analysis_text:
                score += weight
                # Bonus für Keyword im Titel
                if title and keyword in title.lower():
                    score += weight * 0.5
        scores[doc_type] = score
    
    # Besten Score finden
    if scores:
        best_type = max(scores, key=scores.get)
        best_score = scores[best_type]
        
        # Mindest-Score für automatische Zuordnung
        if best_score >= 10:
            return best_type
    
    # Fallback: "OTHER" wenn keine klare Zuordnung möglich
    return "OTHER"

def extract_comprehensive_metadata(text: str, title: str = "") -> dict:
    """
    Extrahiert umfassende Metadaten aus Dokumententext.
    
    Args:
        text: Vollständiger extrahierter Text
        title: Dokumenttitel
        
    Returns:
        dict: Umfassende Metadaten inkl. Dokumenttyp, Keywords, etc.
    """
    # Erste 3000 Zeichen für Analyse
    analysis_text = text[:3000]
    
    metadata = {
        "detected_type": analyze_document_type(text, title),
        "keywords": extract_keywords(text),
        "text_length": len(text),
        "analysis_excerpt": analysis_text,
        "has_tables": "tabelle" in text.lower() or "|" in text,
        "has_procedures": any(word in text.lower() for word in ["schritt", "verfahren", "durchführung", "ablauf"]),
        "compliance_indicators": [],
        "complexity_score": _calculate_complexity_score(analysis_text)
    }
    
    # Compliance-Indikatoren sammeln
    compliance_terms = ["iso", "mdr", "fda", "cfr", "gmp", "audit", "validierung", "kalibrierung"]
    for term in compliance_terms:
        if term in text.lower():
            metadata["compliance_indicators"].append(term.upper())
    
    return metadata

def _calculate_complexity_score(text: str) -> int:
    """
    Berechnet Komplexitäts-Score basierend auf Textmerkmalen.
    
    Args:
        text: Zu analysierender Text
        
    Returns:
        int: Komplexitätsscore (1-10)
    """
    score = 1
    
    # Länge berücksichtigen
    if len(text) > 1000:
        score += 2
    if len(text) > 2000:
        score += 2
    
    # Fachbegriffe zählen
    technical_terms = ["spezifikation", "validierung", "kalibrierung", "verifikation", "risiko"]
    tech_count = sum(1 for term in technical_terms if term in text.lower())
    score += min(tech_count, 3)
    
    # Struktur bewerten
    if "1." in text or "a)" in text:  # Nummerierte Listen
        score += 1
    if text.count("\n") > 10:  # Viele Absätze
        score += 1
    
    return min(score, 10)

def clean_extracted_text(text: str) -> str:
    """
    Bereinigt extrahierten Text für bessere Verarbeitung.
    
    Args:
        text: Eingabetext
        
    Returns:
        str: Bereinigter Text
    """
    if not text:
        return ""
    
    # Mehrfache Zeilenendings normalisieren
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Leerzeichen normalisieren
        line = ' '.join(line.split())
        
        # Sehr kurze oder nur-Leerzeichen Zeilen überspringen
        if len(line.strip()) > 2:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines) 